"""Document loaders for PDF, DOCX, TXT and MD files.

Each loader returns text split into "pages" so that retrieved chunks can later
be cited with a page number. PDFs have real pages (via PyMuPDF); DOCX/TXT/MD
do not, so they are treated as a single page (page 1).

``load_document`` is the public entry point. It returns a dict::

    {
        "source_file": "ICH_Q7.pdf",
        "file_type": "pdf",
        "doc_type": "ICH",            # inferred regulatory category
        "pages": [{"page_number": 1, "text": "..."}, ...],
    }
"""

from __future__ import annotations

import os

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from src.utils.config import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def infer_doc_type(filename: str) -> str:
    """Infer a coarse regulatory document category from the file name."""
    name = filename.lower()
    # Match "ich" or ICH quality-guideline codes like q7, q9, q1a, q3d, ...
    if "ich" in name or any(f"q{d}" in name for d in "123456789"):
        return "ICH"
    if "fda" in name:
        return "FDA"
    if "who" in name:
        return "WHO GMP"
    if "schedule" in name and "m" in name:
        return "CDSCO Schedule M"
    if "cdsco" in name:
        return "CDSCO"
    if "sop" in name:
        return "SOP"
    if "gmp" in name:
        return "GMP"
    return "Regulatory Document"


def _load_pdf(path: str) -> list[dict]:
    """Extract text page by page from a PDF using PyMuPDF."""
    pages: list[dict] = []
    with fitz.open(path) as doc:
        for index, page in enumerate(doc):
            text = page.get_text("text") or ""
            if text.strip():
                pages.append({"page_number": index + 1, "text": text})
    return pages


def _load_docx(path: str) -> list[dict]:
    """Extract text from a DOCX file (paragraphs + tables) as a single page."""
    document = DocxDocument(path)

    parts: list[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Include table cell text, which python-docx does not expose via paragraphs.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return [{"page_number": 1, "text": text}] if text.strip() else []


def _load_text(path: str) -> list[dict]:
    """Read a plain text or markdown file as a single page."""
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        text = handle.read()
    return [{"page_number": 1, "text": text}] if text.strip() else []


def load_document(file_path: str) -> dict:
    """Load a document from disk and return its text split into pages.

    Raises ``ValueError`` for unsupported file extensions.
    """
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    logger.info("Loading document: %s", filename)

    if ext == ".pdf":
        pages = _load_pdf(file_path)
    elif ext == ".docx":
        pages = _load_docx(file_path)
    else:  # .txt or .md
        pages = _load_text(file_path)

    if not pages:
        logger.warning("No extractable text found in %s", filename)

    return {
        "source_file": filename,
        "file_type": ext.lstrip("."),
        "doc_type": infer_doc_type(filename),
        "pages": pages,
    }
